"""
Document text extraction for PDF, DOCX, HTML, and plain text.

Each extractor returns a ``Document`` with the extracted text and metadata.
Format is auto-detected from file extension or MIME type.

All extraction libraries are optional dependencies::

    uv add ai-toolkit[ingestion]   # pypdf, python-docx, beautifulsoup4
"""

from __future__ import annotations

import mimetypes
import os
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any


class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TEXT = "text"
    MARKDOWN = "markdown"


@dataclass
class Document:
    """Extracted document with text content and metadata."""

    content: str
    """Full extracted text."""

    source: str
    """Original file path, URL, or identifier."""

    format: DocumentFormat
    """Detected document format."""

    page_count: int = 0
    """Number of pages (PDF only, 0 for other formats)."""

    char_count: int = 0
    """Total character count of extracted text."""

    metadata: dict[str, Any] = field(default_factory=dict)
    """Additional metadata (author, title, dates, etc.)."""

    def __post_init__(self) -> None:
        self.char_count = len(self.content)


# ─── Format Detection ────────────────────────────────────────────────────────

_EXT_MAP: dict[str, DocumentFormat] = {
    ".pdf": DocumentFormat.PDF,
    ".docx": DocumentFormat.DOCX,
    ".doc": DocumentFormat.DOCX,
    ".html": DocumentFormat.HTML,
    ".htm": DocumentFormat.HTML,
    ".txt": DocumentFormat.TEXT,
    ".md": DocumentFormat.MARKDOWN,
    ".markdown": DocumentFormat.MARKDOWN,
    ".csv": DocumentFormat.TEXT,
    ".json": DocumentFormat.TEXT,
    ".xml": DocumentFormat.TEXT,
}

_MIME_MAP: dict[str, DocumentFormat] = {
    "application/pdf": DocumentFormat.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentFormat.DOCX,
    "application/msword": DocumentFormat.DOCX,
    "text/html": DocumentFormat.HTML,
    "text/plain": DocumentFormat.TEXT,
    "text/markdown": DocumentFormat.MARKDOWN,
}


def detect_format(
    path: str | None = None,
    *,
    mime_type: str | None = None,
) -> DocumentFormat:
    """
    Detect document format from file extension or MIME type.

    Raises ValueError if format cannot be determined.
    """
    if mime_type:
        fmt = _MIME_MAP.get(mime_type)
        if fmt:
            return fmt

    if path:
        ext = Path(path).suffix.lower()
        fmt = _EXT_MAP.get(ext)
        if fmt:
            return fmt

        # Try mimetypes module
        guessed, _ = mimetypes.guess_type(path)
        if guessed:
            fmt = _MIME_MAP.get(guessed)
            if fmt:
                return fmt

    raise ValueError(
        f"Cannot detect format for path={path!r}, mime_type={mime_type!r}. "
        f"Supported: {', '.join(f.value for f in DocumentFormat)}"
    )


# ─── Extractors ──────────────────────────────────────────────────────────────


def extract_pdf(source: str | bytes, *, path: str = "<bytes>") -> Document:
    """
    Extract text from a PDF file.

    Args:
        source: File path (str) or raw PDF bytes
        path: Label for the source (used when source is bytes)
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "pypdf is required for PDF extraction. Run: uv add pypdf"
        ) from e

    import io

    if isinstance(source, str):
        path = source
        reader = PdfReader(source)
    else:
        reader = PdfReader(io.BytesIO(source))

    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    metadata: dict[str, Any] = {}
    if reader.metadata:
        for key in ("title", "author", "subject", "creator"):
            val = getattr(reader.metadata, key, None)
            if val:
                metadata[key] = val

    return Document(
        content="\n\n".join(pages),
        source=path,
        format=DocumentFormat.PDF,
        page_count=len(reader.pages),
        metadata=metadata,
    )


def extract_docx(source: str | bytes, *, path: str = "<bytes>") -> Document:
    """
    Extract text from a DOCX file.

    Args:
        source: File path (str) or raw DOCX bytes
        path: Label for the source (used when source is bytes)
    """
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction. Run: uv add python-docx"
        ) from e

    import io

    if isinstance(source, str):
        path = source
        doc = docx.Document(source)
    else:
        doc = docx.Document(io.BytesIO(source))

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract tables too
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    metadata: dict[str, Any] = {}
    if doc.core_properties:
        for key in ("title", "author", "subject", "created", "modified"):
            val = getattr(doc.core_properties, key, None)
            if val:
                metadata[key] = str(val)

    return Document(
        content="\n\n".join(paragraphs),
        source=path,
        format=DocumentFormat.DOCX,
        metadata=metadata,
    )


def extract_html(source: str | bytes, *, path: str = "<bytes>") -> Document:
    """
    Extract text from HTML, stripping tags.

    Args:
        source: File path (str), raw HTML bytes, or HTML string.
            If source is a str that looks like a file path, reads the file.
            Otherwise treats it as raw HTML.
        path: Label for the source
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError(
            "beautifulsoup4 is required for HTML extraction. Run: uv add beautifulsoup4"
        ) from e

    if isinstance(source, bytes):
        html_text = source.decode("utf-8", errors="replace")
    elif os.path.isfile(source):
        path = source
        with open(source, encoding="utf-8", errors="replace") as f:
            html_text = f.read()
    else:
        # Treat as raw HTML string
        html_text = source

    soup = BeautifulSoup(html_text, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    metadata: dict[str, Any] = {}
    title_tag = soup.find("title")
    if title_tag and title_tag.string:
        metadata["title"] = title_tag.string.strip()

    return Document(
        content=text,
        source=path,
        format=DocumentFormat.HTML,
        metadata=metadata,
    )


def extract_text(source: str | bytes, *, path: str = "<bytes>") -> Document:
    """
    Extract text from a plain text file (or markdown, CSV, JSON, etc.).

    Args:
        source: File path (str) or raw bytes/text
        path: Label for the source
    """
    if isinstance(source, bytes):
        content = source.decode("utf-8", errors="replace")
    elif os.path.isfile(source):
        path = source
        with open(source, encoding="utf-8", errors="replace") as f:
            content = f.read()
    else:
        # Treat as raw text string
        content = source

    fmt = DocumentFormat.TEXT
    if path.endswith((".md", ".markdown")):
        fmt = DocumentFormat.MARKDOWN

    return Document(
        content=content,
        source=path,
        format=fmt,
    )


# ─── Auto-detect Extraction ─────────────────────────────────────────────────

_EXTRACTOR_MAP = {
    DocumentFormat.PDF: extract_pdf,
    DocumentFormat.DOCX: extract_docx,
    DocumentFormat.HTML: extract_html,
    DocumentFormat.TEXT: extract_text,
    DocumentFormat.MARKDOWN: extract_text,
}


def extract_document(
    source: str | bytes,
    *,
    path: str | None = None,
    format: DocumentFormat | None = None,
    mime_type: str | None = None,
) -> Document:
    """
    Extract text from a document, auto-detecting format.

    Args:
        source: File path (str) or raw bytes
        path: Label/path for format detection (required when source is bytes)
        format: Explicit format override (skips detection)
        mime_type: MIME type for format detection

    Returns:
        Document with extracted text and metadata
    """
    if format is None:
        detect_path = path if isinstance(source, bytes) else source if isinstance(source, str) else None
        format = detect_format(detect_path, mime_type=mime_type)

    extractor = _EXTRACTOR_MAP.get(format)
    if not extractor:
        raise ValueError(f"No extractor for format: {format}")

    label = path or (source if isinstance(source, str) else "<bytes>")
    return extractor(source, path=label)
